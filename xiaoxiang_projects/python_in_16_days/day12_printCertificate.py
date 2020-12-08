# Practice python in 16 days
# https://clock.kaiwenba.com/clock/course?sn=ExhvY&course_id=16

'''
  判断当前目录下是否有‘录取同学名单.xlsx’和‘录取通知书模板.docx’
  自动从excel表格提取信息后打印录取通知书生成 pdf文件
'''

#os.listdir() 方法用于返回指定的文件夹包含的文件或文件夹名字的列表。
from os import listdir
xlsx_name=''
docx_name=''
for file in listdir():
    if '录取同学名单.xlsx' in file:
        xlsx_name = file
        print("录取学生信息存放在'"+xlsx_name+"'中")
    if '录取通知书模板.docx' in file:
        docx_name = file
        print("想要修改的模板是'"+docx_name+"'")


# 对于Excel的操作主要是从表格中读取学生的数据信息。
# 这里我们要用到神奇的Openpyxl库。Openpyxl是一个读写Excel的Python库，是一款比较综合的工具。
# 不仅能够读取和修改Excel文档，而且可以对Excel文件内单元格进行详细设置，
# 包括单元格样式等内容，甚至还支持图表插入、打印设置等内容。
# 使用它也可以处理数据量较大的Excel文件。
# 你需要先掌握三个对象：
# Workbook(工作簿，一个包含多个Sheet的Excel文件)
# Worksheet（工作表，一个Workbook有多个Worksheet，如“Sheet1”,“Sheet2”等）
# Cell（单元格，存储具体的数据对象）

# 首先打开Excel文件’录取同学名单.xlsx’。
from openpyxl import load_workbook
wb = load_workbook(xlsx_name)

# 下面用sheetx0得到Excel文件中的所有工作表(案例中Excel文件只有一个工作表)
sheetx0 = wb.sheetnames

# 选择我们所要操作的工作表。
sheetx = wb[sheetx0[0]]

# 来查看一下Excel表格中共有多少行数据。
print(sheetx.max_row) #18

# 接下来可以按照行和列读取Excel单元格的数据，来看下第一行第一列的数据吧～
print("第一行第一列的数据",sheetx.cell(row=1,column=1).value)

# 想要获取Excel中第一行的所有数据该怎么做呢？只需要一个循环就能轻松实现了。
print("sheetx.max_column",sheetx.max_column)
print("第一行的所有数据")
for l in range(1,sheetx.max_column+1):
    print("第{}列数据".format(l),sheetx.cell(row=1,column=l).value)

# 对Word文档操作用到docx库。
# another method: docx2txt:http://theautomatic.net/2019/10/14/how-to-read-word-documents-with-python/
# docx库中的方法Document可以对Word文档进行读写操作。
# 首先打开Word文件，读取内容的所有段落。

from docx import Document
#打开Word文档
document = Document(docx_name)
#读取所有的自然段
all_paragraphs = document.paragraphs

# 下面我们自定义一个方法，使它能够实现新旧文本的替换的功能：
def replace_text(old_text, new_text):
    #读取所有的自然段
    all_paragraphs = document.paragraphs
    for paragraph in all_paragraphs:
        for run in paragraph.runs:
            run_text = run.text.replace(old_text, new_text)
            run.text = run_text

# 生成并保存通知书文件
import os
# 表格中的数据共有18行，第一行是占位符，第二行是表头，从第三行开始是每位同学的入学信息。
# 所以在进行操作时，需要从第三行数据开始读取，把表中内容和Word文本中占位符替换，生成录取通知书文件。
# 后面继续对每行进行相同操作就能得到同学们的录取通知书文档了。
for row in range(3,sheetx.max_row+1):
    document = Document(docx_name)
    # openpyxl在使用sheetx.max_column时可能会读取到空的单元格，这里进行剔除
    if sheetx.cell(row=row,column=1).value!=None:
        # 录取通知书要素Excel中逐行循环
        for l in range(1,sheetx.max_column+1):
            # print(l,row)
            # 获取Excel第一行的内容
            old_text = sheetx.cell(row=1,column=l).value
            # print("old_text", old_text)
            # 对循环的Excel当前列逐行读取新要素
            new_text = sheetx.cell(row=row,column=l).value
            # print("new_text",new_text)
            # 进行替换
            replace_text(str(old_text),str(new_text))
            # 定义文件名为当前列第一行的内容
            filename = str(sheetx.cell(row=row,column=1).value)
            # print("filename",filename)
        # 按定义的文件名进行保存，此处的文件夹是需要提前手动建立好的
        document.save(r"存放学生录取通知书Word文档\\%s.docx"%(filename))
        print('学号{}学生的录取通知书已生成'.format(filename))

# 把通知书批量转换为pdf格式。
# 完成转换操作需要下面三个步骤：
# 启动Word进程
# 遍历Word所在目录下的所有Word文件一一转化pdf
# 保存pdf
# 这里用到的模块是win32，它不是纯粹的Python模块，是windows功能和Office功能来实现的。
# import win32com.client #pip install pywin32
# wordApp = win32com.client.Dispatch('word.Application')

# 控制Word转化pdf过程是否可视化, 这里选择不用可视化
# wordApp.Visible = False

# 用其中一个文件转换一下实验：
# # 接着先打开一个学生录取通知书的Word文档：
# import os
# filepath = os.getcwd()
# filepath = filepath+'/'+'存放学生录取通知书Word文档/20190601.docx'
# myDoc = wordApp.Documents.Open(filepath)
# print("myDoc",myDoc)
#
# # 打开之后将它转化为pdf格式：
# import os
# output_file_path = os.getcwd()   # PDF储存位置
# #此处的文件夹"存放学生录取通知书PDF文档"是需要提前手动建立好的
# output_file_path = output_file_path + "/" + "存放学生录取通知书PDF文档/20190601.pdf"
# myDoc.ExportAsFixedFormat(output_file_path, 17, Item=7, CreateBookmarks=0)
# # 可以看到成功生成了pdf格式的文件，最后关闭word进程：
# wordApp.Quit()

# 接下来我们把上面的转换pdf的步骤整合在一起：
class Word_2_PDF(object):
    def __init__(self, filepath, wordVisible = False):
        """
        :param filepath:
        :param Debug: 控制过程是否可视化
        """
        # 启动独立的进程
        self.wordApp = win32com.client.Dispatch('word.Application')
        # 控制过程是否可视化
        self.wordApp.Visible = wordVisible
        # 打开文件夹
        self.myDoc = self.wordApp.Documents.Open(filepath)
    def export_pdf(self, output_file_path):
        """
        将Word文档转化为PDF文件
        :param output_file_path:
        :return:
        """
        self.myDoc.ExportAsFixedFormat(output_file_path, 17, Item=7, CreateBookmarks=0)
    def close(self):
        self.wordApp.Quit()


import win32com.client
import pythoncom
import os
rootpath = os.getcwd()  # 文件夹路径
save_path = os.getcwd()   # PDF储存位置
if not os.path.exists('存放学生录取通知书PDF文档'):
    os.mkdir('存放学生录取通知书PDF文档')
rootpath = rootpath + "\\" + "存放学生录取通知书Word文档"
save_path = save_path + "\\" + "存放学生录取通知书PDF文档"

# 遍历目录下所有的Word文件，用同样的方法进行转换就没问题啦。
pythoncom.CoInitialize()
# os_dict = {root: [dirs, files] for root, dirs, files in os.walk(rootpath)}
for parent, dirnames, filenames in os.walk(rootpath):
    for filename in filenames:
        if u'.doc' in filename and u'~$' not in filename:
            # 直接保存为PDF文件
            print(rootpath+filename)
            a = Word_2_PDF(rootpath + '\\' + filename, True)
            title = filename.split('.')[0]  # 删除.docx
            a.export_pdf(save_path + '\\' + title + '.pdf')
            a.close()
print('转化完成')