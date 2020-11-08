# Practice python in 16 days
# https://clock.kaiwenba.com/clock/course?sn=ExhvY&course_id=71

'''
 按模版批量生成Word文件（下）
 自动替换信息，批量生成邀请函文件。
'''
from docx import Document
from openpyxl import load_workbook

#打开word文档
document = Document('C:/Users/Maggie/Desktop/Udemy-desktop/xiaoxiangxueyuan/Python认知打卡课课程资料包/资料代码整理/四五关代码/第五关代码/邀请函模版.docx')

#读取所有段落
all_paragraphs = document.paragraphs

#读取excel表格
wb = load_workbook('C:/Users/Maggie/Desktop/Udemy-desktop/xiaoxiangxueyuan/Python认知打卡课课程资料包/资料代码整理/四五关代码/第五关代码/file_merge.xlsx')
sheet = wb.active
sheet.max_row

def replace_text(old_text, new_text):
    #读取所有段落
    all_paragraphs = document.paragraphs
    for paragraph in all_paragraphs:
        for run in paragraph.runs:
            run_text = run.text.replace(old_text, new_text)
            run.text = run_text

for row in range(2,sheet.max_row+1):
    # document = Document('/Users/cuiyingdan/Desktop/第五关/邀请函模版.docx')
    old_text='×××'
    new_text=sheet.cell(row=row,column=1).value
    replace_text(str(old_text),str(new_text))
    filename=str(sheet.cell(row=row,column=1).value)
    document.save('{}.docx'.format(filename))